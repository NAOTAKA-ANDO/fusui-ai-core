import os
from pathlib import Path
from docx import Document
import json

# 処理対象の基点フォルダ
BASE_DIR = Path("E:/風水AI開発/knowledge")  # ご環境に合わせて絶対パス指定

def extract_docx_structure(docx_path: Path, rel_root: Path):
    doc = Document(docx_path)
    full_text = ""
    images = []
    img_counter = 1

    # 相対フォルダ名（書籍名）取得
    book_name = rel_root.name
    output_json_dir = BASE_DIR.parent / "structured" / book_name / "chapters_json"
    output_img_dir = BASE_DIR.parent / "structured" / book_name / "images"
    output_json_dir.mkdir(parents=True, exist_ok=True)
    output_img_dir.mkdir(parents=True, exist_ok=True)

    # 画像抽出
    for rel in doc.part._rels:
        rel_obj = doc.part._rels[rel]
        if "image" in rel_obj.target_ref:
            img_data = rel_obj.target_part.blob
            img_filename = f"{docx_path.stem}_img{img_counter}.png"
            img_path = output_img_dir / img_filename
            with open(img_path, "wb") as img_file:
                img_file.write(img_data)
            images.append(str(img_path.name))
            full_text += f"\n[image{img_counter}]\n"
            img_counter += 1

    # 本文抽出
    for para in doc.paragraphs:
        full_text += para.text + "\n"

    # 構造化
    result = {
        "filename": docx_path.name,
        "title": docx_path.stem,
        "word_count": len(full_text),
        "paragraphs": len(doc.paragraphs),
        "content": full_text.strip(),
        "images": images
    }

    # JSON保存
    json_path = output_json_dir / f"{docx_path.stem}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"✅ 完了：{docx_path.name} → {json_path.name}")

# 再帰的にchaptersフォルダを探す
for root, dirs, files in os.walk(BASE_DIR):
    if Path(root).name == "chapters":
        chapter_dir = Path(root)
        book_root = chapter_dir.parent
        for docx_file in chapter_dir.glob("*.docx"):
            extract_docx_structure(docx_file, book_root)

